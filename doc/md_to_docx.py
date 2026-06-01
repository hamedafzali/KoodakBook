#!/usr/bin/env python3
"""Converts project.md to KoodakBook_ProductVision.docx. Run from any directory."""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
MD_FILE = ROOT / "project.md"
OUT_FILE = Path(__file__).resolve().parent / "KoodakBook_ProductVision.docx"


def add_table_from_md(doc, lines):
    rows = [l for l in lines if l.startswith("|")]
    if len(rows) < 2:
        return
    # filter separator row
    rows = [r for r in rows if not re.match(r"^\|[-| :]+\|$", r.strip())]
    if not rows:
        return

    cols = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = col
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for row_text in rows[1:]:
        cells = [c.strip() for c in row_text.strip().strip("|").split("|")]
        row = table.add_row().cells
        for i, cell in enumerate(cells):
            if i < len(row):
                row[i].text = cell

    doc.add_paragraph()


def clean(text):
    # remove bold/italic markers
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def convert():
    if not MD_FILE.exists():
        print(f"project.md not found at {MD_FILE}")
        sys.exit(1)

    doc = Document()

    # document title
    title_para = doc.add_heading("KoodakBook — Product Vision Document", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Helping diaspora children learn Persian through stories, play, and family connection."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lines = MD_FILE.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # skip the h1 title and blockquote (already added above)
        if line.startswith("# ") or line.startswith("> "):
            i += 1
            continue

        # horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # headings
        if line.startswith("#### "):
            doc.add_heading(clean(line[5:]), 4)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(clean(line[4:]), 3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(clean(line[3:]), 1)
            i += 1
            continue

        # table — collect all consecutive table lines
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph("\n".join(code_lines))
            p.style = "No Spacing"
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            doc.add_paragraph()
            continue

        # bullet lists (-, *, [ ])
        bullet_match = re.match(r"^(\s*)[-*] (.+)", line)
        checkbox_match = re.match(r"^(\s*)- \[[ x]\] (.+)", line)
        if checkbox_match:
            indent = len(checkbox_match.group(1)) // 2
            style = "List Bullet 2" if indent > 0 else "List Bullet"
            p = doc.add_paragraph(style=style)
            checked = "x" in line[line.index("[") + 1]
            prefix = "[x] " if checked else "[ ] "
            p.add_run(prefix + clean(checkbox_match.group(2)))
            i += 1
            continue
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            style = "List Bullet 2" if indent > 0 else "List Bullet"
            text = bullet_match.group(2)
            p = doc.add_paragraph(style=style)
            bold_match = re.match(r"\*\*(.+?)\*\*:? ?(.*)", text)
            if bold_match:
                run = p.add_run(bold_match.group(1))
                run.bold = True
                rest = bold_match.group(2)
                if rest:
                    p.add_run(": " + clean(rest))
            else:
                p.add_run(clean(text))
            i += 1
            continue

        # numbered lists
        num_match = re.match(r"^\d+\. (.+)", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            p.add_run(clean(num_match.group(1)))
            i += 1
            continue

        # bold line acting as a sub-header
        if re.match(r"^\*\*.+\*\*$", line.strip()):
            p = doc.add_paragraph()
            run = p.add_run(clean(line.strip()))
            run.bold = True
            i += 1
            continue

        # empty line
        if line.strip() == "":
            i += 1
            continue

        # normal paragraph
        doc.add_paragraph(clean(line))
        i += 1

    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    convert()
